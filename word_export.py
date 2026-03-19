"""Exports transcript and summary to a Word document, returned as bytes."""
import io
import re

from docx import Document
from docx.shared import Pt


def _sanitize(text: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)


def _add_body(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(_sanitize(text))
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    return p


def _add_header(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(_sanitize(text))
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    return p


def export_to_word(title: str, clean_text: str, summary_text: str) -> bytes:
    """
    Builds a Word document and returns it as bytes (no disk I/O).
    Body text: Calibri 12pt
    Section headers: Calibri 12pt bold
    Exec summary timestamps: bold
    """
    doc = Document()

    # Set Normal style defaults
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(_sanitize(title))
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)
    title_run.bold = True

    doc.add_paragraph()

    # Executive Summary
    _add_header(doc, "Executive Summary")
    doc.add_paragraph()

    ts_pattern = re.compile(r"(\[\d{1,2}:\d{2}\])")
    for line in summary_text.splitlines():
        if not line.strip():
            continue
        p = doc.add_paragraph()
        for part in ts_pattern.split(_sanitize(line.strip())):
            run = p.add_run(part)
            run.font.name = "Calibri"
            run.font.size = Pt(12)
            if ts_pattern.match(part):
                run.bold = True

    doc.add_page_break()

    # Full Transcript
    _add_header(doc, "Full Transcript")
    doc.add_paragraph()

    for line in clean_text.splitlines():
        if line.startswith("## "):
            doc.add_paragraph()
            _add_header(doc, line[3:].strip())
        elif line.strip():
            _add_body(doc, line.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
